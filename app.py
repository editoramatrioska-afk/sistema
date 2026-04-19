import streamlit as st
from supabase import create_client, Client
import math
from docx import Document
from fpdf import FPDF
import io
import os 
from lxml import etree
from datetime import datetime
from num2words import num2words
import pandas as pd

# --- 1. CONFIGURAÇÃO E CONEXÃO ---
st.set_page_config(page_title="Editora Matrioska - Sistema Interno", layout="wide")

URL: str = "https://gbeoizrqxzopjsxthwym.supabase.co"
KEY: str = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6ImdiZW9penJxeHpvcGpzeHRod3ltIiwicm9sZSI6ImFub24iLCJpYXQiOjE3NzY0NzAwNzcsImV4cCI6MjA5MjA0NjA3N30.dGQ3gnzjT5jHd4LAZTTSp1k8XemowUglFToPbDL38OY"
supabase: Client = create_client(URL, KEY)

NOME_LOGO = "logo.png"
NOME_RODAPE = "rodape.png"

# --- 2. FUNÇÕES DE APOIO E AUTENTICAÇÃO ---

def login():
    st.container()
    col_l, col_c, col_r = st.columns([1, 2, 1])
    with col_c:
        if os.path.exists(NOME_LOGO): 
            st.image(NOME_LOGO, width=200)
        st.title("🔐 Acesso Restrito")
        
        with st.form("login_form"):
            email = st.text_input("E-mail")
            password = st.text_input("Senha", type="password")
            submit = st.form_submit_button("Entrar", use_container_width=True)
            
            if submit:
                try:
                    res = supabase.auth.sign_in_with_password({"email": email, "password": password})
                    st.session_state['logged_in'] = True
                    st.session_state['user_email'] = email
                    st.success("Autenticado! Entrando...")
                    st.rerun()
                except Exception:
                    st.error("Usuário ou senha inválidos.")

def valor_por_extenso(valor):
    try:
        inteiro = int(valor)
        extenso = num2words(inteiro, lang='pt_BR')
        return f"({extenso} reais)"
    except:
        return ""

def contar_caracteres_oficial_word(arquivo):
    doc = Document(arquivo)
    texto_total = []
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    for p in doc.paragraphs: texto_total.append(p.text)
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells: texto_total.append(celula.text)
    for rel in doc.part.rels.values():
        if "footnotes" in rel.target_ref or "endnotes" in rel.target_ref:
            root = etree.fromstring(rel.target_part.blob)
            for t in root.xpath('//w:t', namespaces=ns):
                if t.text: texto_total.append(t.text)
    return len("".join(texto_total))

class PDF_Proposta(FPDF):
    def header(self):
        if os.path.exists(NOME_LOGO):
            self.image(NOME_LOGO, (210 - 40) / 2, 8, 40)
        self.ln(25)

    def footer(self):
        if os.path.exists(NOME_RODAPE):
            self.image(NOME_RODAPE, 0, 275, 210)

def obter_data_formatada():
    meses = ["janeiro", "fevereiro", "março", "abril", "maio", "junho", 
             "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"]
    hoje = datetime.now()
    return f"São Paulo, {hoje.day} de {meses[hoje.month - 1]} de {hoje.year}"

def gerar_pdf_matrioska(dados):
    pdf = PDF_Proposta()
    pdf.set_auto_page_break(auto=True, margin=35)
    
    # --- PÁGINA 1: APRESENTAÇÃO ---
    pdf.add_page()
    pdf.set_font("helvetica", 'B', 12)
    pdf.cell(0, 10, obter_data_formatada(), ln=True)
    pdf.ln(5)
    pdf.set_font("helvetica", size=12)
    texto_apresentacao = (
        "A Matrioska é uma editora de livros científicos, técnicos e profissionais, provedora de conteúdo para a "
        "formação sólida de estudantes universitários e para a atualização de profissionais das mais diversas áreas do conhecimento.\n\n"
        "Nosso catálogo contempla autoras e autores de primeira linha, que aliam o que há de mais inovador em termos de abordagem "
        "e rigor acadêmico, avaliados e validados por um renomado Conselho Editorial Nacional e Internacional que resguarda a "
        "qualidade de nossas publicações.\n\n"
        "Se você acredita no futuro dos livros (independente do formato ou suporte), vem com a gente!\n"
        "Para nós é uma grande satisfação tê-la como nosso autora.\n\n"
        "Forte abraço,\n\n"
        "Patrícia Melo e Luciana Félix"
    )
    pdf.multi_cell(0, 7, texto_apresentacao)

    # --- PÁGINA 2: PROJETO ---
    pdf.add_page()
    pdf.set_font("helvetica", 'B', 14)
    pdf.cell(0, 10, "Projeto Editorial", ln=True)
    pdf.ln(5)
    pdf.set_font("helvetica", size=12)
    pdf.cell(0, 8, f"Livro: {dados['livro']}", ln=True)
    pdf.cell(0, 8, f"Autor: {dados['cliente']}", ln=True)
    pdf.ln(2)
    pdf.set_font("helvetica", 'B', 12)
    pdf.cell(0, 8, "Especificações:", ln=True)
    pdf.set_font("helvetica", size=12)
    pdf.cell(0, 7, f"- Laudas: {dados['laudas']:.0f}", ln=True)
    pdf.cell(0, 7, f"- Páginas estimadas: {dados['paginas']}", ln=True)
    pdf.cell(0, 7, f"- Formato: {dados['formato']}", ln=True)
    pdf.cell(0, 7, f"- Miolo: {dados['miolo']}", ln=True)
    pdf.cell(0, 7, f"- Capa: {dados['capa']}", ln=True)
    pdf.cell(0, 7, f"- Acabamento: {dados['acabamento']}", ln=True)
    pdf.ln(5)

    pdf.set_font("helvetica", 'B', 12)
    pdf.cell(0, 8, "Produção editorial Premium", ln=True)
    pdf.set_font("helvetica", size=11)
    itens = [
        "- Copidesque e preparação de textos (revisão ortográfica e padronização conforme ABNT), diagramação, revisão pós-diagramação, conferências e fechamento de arquivo.",
        "- ISBN, Capa; ficha catalográfica, código de barras.",
        "- Conteúdo publicado sob o selo editorial Matrioska Editora.",
        "- Edições impressa e digital poderão ser disponibilizadas na loja virtual da editora e nas principais plataformas de e-commerce e livrarias virtuais (Amazon e marketplaces)."
    ]
    for item in itens:
        pdf.multi_cell(0, 6, item)
        pdf.ln(2)

    # --- PÁGINA 3: INVESTIMENTO ---
    pdf.add_page()
    pdf.set_font("helvetica", 'B', 14)
    pdf.cell(0, 10, "Proposta de investimento:", ln=True)
    pdf.ln(5)
    pdf.set_font("helvetica", size=12)
    valor_f = f"R$ {dados['total']:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    extenso = valor_por_extenso(dados['total'])
    texto_investimento = (
        f"- {valor_f} {extenso} para custeio da produção editorial "
        "(etapas de copidesque, projeto gráfico e diagramação, revisão pós-diagramação, "
        "capa, conferências e fechamento de arquivos: para impressão e e-books);\n\n"
        "- Não inclui exemplares impressos;\n\n"
        "- Condição de pagamento: 40% na assinatura do contrato; 30% após 30 dias "
        "e o restante no envio do arquivo para a gráfica.\n\n"
        "Orçamento válido por 30 dias."
    )
    pdf.multi_cell(0, 7, texto_investimento)
    pdf.ln(20)
    pdf.set_font("helvetica", 'B', 12)
    pdf.cell(0, 10, obter_data_formatada(), ln=True)
    return bytes(pdf.output())

# --- 3. INICIALIZAÇÃO DE ESTADO ---
if 'logged_in' not in st.session_state: st.session_state['logged_in'] = False
if 'edit_data' not in st.session_state: st.session_state['edit_data'] = None

# --- 4. LÓGICA DE NAVEGAÇÃO ---
if not st.session_state['logged_in']:
    login()
else:
    tab1, tab2 = st.tabs(["📝 Novo Orçamento", "📜 Histórico"])

    # --- ABA 1: GERADOR ---
    with tab1:
        with st.sidebar:
            if os.path.exists(NOME_LOGO): st.image(NOME_LOGO, width=150)
            st.write(f"Logado como: {st.session_state['user_email']}")
            st.header("Configuração de Preços")
            p_copidesque = st.number_input("Copidesque (lauda):", value=6.0)
            p_diagramacao = st.number_input("Diagramação (pág):", value=5.0)
            p_revisao = st.number_input("Revisão (pág):", value=4.0)
            p_capa = st.number_input("Capa:", value=500.0)
            p_isbn = st.number_input("ISBN/Fichas:", value=300.0)
            p_epub = st.number_input("E-pub (pág):", value=2.50)
            p_taxa_editora = st.number_input("Taxa Editora:", value=3000.0)
            if st.button("Sair"):
                st.session_state['logged_in'] = False
                st.rerun()

        st.title("📚 Sistema de Orçamento - Matrioska")
        
        edit = st.session_state['edit_data']
        col1, col2 = st.columns(2)
        
        with col1:
            nome_cliente = st.text_input("Autor(a):", value=edit['cliente'] if edit else "")
            nome_livro = st.text_input("Nome do Livro:", value=edit['livro'] if edit else "")
            arquivo_word = st.file_uploader("Subir Manuscrito (.docx)", type=["docx"])
            
            if arquivo_word:
                total_caracteres = contar_caracteres_oficial_word(arquivo_word)
                st.success(f"{total_caracteres} caracteres detectados.")
            else:
                total_caracteres = st.number_input("Total caracteres manualmente:", value=edit['caracteres'] if edit else 0)
            
            pag_extras = st.number_input("Páginas extras:", min_value=0, value=0)

        with col2:
            formato_sel = st.selectbox("Formato:", ["14x21", "16x23", "17x24"], 
                                      index=["14x21", "16x23", "17x24"].index(edit['formato']) if edit and edit['formato'] in ["14x21", "16x23", "17x24"] else 0)
            miolo_sel = st.selectbox("Miolo:", ["PB", "Colorido", "PB com caderno colorido"])
            capa_sel = st.selectbox("Capa:", ["4x0", "4x1", "4x4"])
            acabamento_sel = st.selectbox("Acabamento:", ["Brochura", "Capa Dura"])
            incluir_epub = st.checkbox("Incluir E-pub?", value=True)

        meta = {"14x21": 1700, "16x23": 2200, "17x24": 2900}.get(formato_sel, 1500)
        qtd_laudas = total_caracteres / 2000
        total_paginas = math.ceil(total_caracteres / meta) + pag_extras if meta > 0 else pag_extras
        
        valor_total = (qtd_laudas * p_copidesque) + (total_paginas * p_diagramacao) + \
                      (total_paginas * p_revisao) + (total_paginas * p_epub if incluir_epub else 0) + \
                      p_capa + p_isbn + p_taxa_editora

        st.markdown("---")
        if total_caracteres > 0:
            st.metric("Investimento Total", f"R$ {valor_total:,.2f}")
            col_b1, col_b2 = st.columns(2)
            
            with col_b1:
                if st.button("💾 Salvar Orçamento"):
                    payload = {
                        "cliente": nome_cliente, "livro": nome_livro, "caracteres": total_caracteres,
                        "formato": formato_sel, "miolo": miolo_sel, "capa": capa_sel,
                        "acabamento": acabamento_sel, "paginas": total_paginas, "valor_total": valor_total
                    }
                    supabase.table("orcamentos").insert(payload).execute()
                    st.success("Salvo no banco de dados!")
            
            with col_b2:
                dados_finais = {"cliente": nome_cliente, "livro": nome_livro, "laudas": qtd_laudas, "formato": formato_sel, 
                               "miolo": miolo_sel, "capa": capa_sel, "acabamento": acabamento_sel, "paginas": total_paginas, "total": valor_total}
                try:
                    pdf_b = gerar_pdf_matrioska(dados_finais)
                    st.download_button("📥 Baixar PDF", pdf_b, f"Proposta_{nome_livro}.pdf", "application/pdf")
                except Exception as e: st.error(f"Erro PDF: {e}")

    # --- ABA 2: HISTÓRICO ---
    with tab2:
        st.title("📜 Histórico de Orçamentos")
        try:
            response = supabase.table("orcamentos").select("*").order("created_at", desc=True).execute()
            if response.data:
                df = pd.DataFrame(response.data)
                cols_hist = ['id', 'created_at', 'cliente', 'livro', 'valor_total']
                cols_exis = [c for c in cols_hist if c in df.columns]
                st.dataframe(df[cols_exis], use_container_width=True)
                
                id_edit = st.selectbox("ID para carregar:", df['id'].tolist())
                if st.button("📂 Carregar Dados"):
                    st.session_state['edit_data'] = next(x for x in response.data if x['id'] == id_edit)
                    st.success("Carregado! Volte na Aba 1.")
            else: st.info("Nenhum registro.")
        except Exception as e: st.error(f"Erro banco: {e}")
